#!/usr/bin/env python3
"""
analisa_edital.py — Análise automática de edital via PNCP + preenchimento Notion

Uso:
  python analisa_edital.py <URL_PNCP> <NOTION_PAGE_ID>

  URL_PNCP pode ser:
    https://pncp.gov.br/app/editais/28645794000160/2026/000054
    ou formato curto: 28645794000160/2026/54

Requer (variáveis de ambiente):
  ANTHROPIC_API_KEY  — chave da API Claude
  NOTION_TOKEN       — token de integração Notion (Settings → Connections → Develop)

Dependências:
  pip install requests pdfplumber anthropic python-dotenv python-docx
"""

import io, json, os, re, sys, zipfile
import requests
import pdfplumber
import docx as docx_lib
import anthropic
from dotenv import load_dotenv

load_dotenv()

# ─── CONFIG ───────────────────────────────────────────────────────────────────
CLAUDE_MODEL   = "claude-sonnet-4-6"
PNCP_BASE      = "https://pncp.gov.br/api/pncp/v1/orgaos"
NOTION_BASE    = "https://api.notion.com/v1"
NOTION_VERSION = "2022-06-28"
MAX_TEXT_CHARS = 60_000

# ─── PNCP ─────────────────────────────────────────────────────────────────────

def parse_pncp(arg: str):
    """Extrai (cnpj, ano, seq) de URL PNCP ou 'cnpj/ano/seq'."""
    clean = arg.replace(".", "").replace("-", "").replace("/", " ").replace("\\", " ")
    m = re.search(r"(\d{14})\D+(\d{4})\D+(\d+)", clean)
    if not m:
        m = re.search(r"(\d{14})/(\d{4})/(\d+)", arg)
    if not m:
        raise ValueError(f"Formato inválido: {arg!r}. Esperado URL PNCP ou cnpj/ano/seq.")
    return m.group(1), int(m.group(2)), int(m.group(3))

def pncp_get(path: str):
    r = requests.get(f"{PNCP_BASE}/{path}", timeout=30)
    r.raise_for_status()
    return r.json()

class ExtracaoInsuficiente(RuntimeError):
    """Levantado quando os documentos não puderam ser lidos com confiança —
    processo deve parar aqui, nunca seguir pro Claude sem base documental real
    (Claude sem grounding inventa habilitação/prazo/critérios)."""


LIMIAR_CHARS_CONFIAVEIS = 500  # abaixo disso, não confia o suficiente pra analisar


def _extrair_pdf(nome: str, content: bytes) -> tuple[bool, str]:
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            # RelacaoItens: prioriza extract_tables para preservar medidas da tabela
            if "relacao" in nome.lower() or "itens" in nome.lower():
                linhas = []
                for pg in pdf.pages:
                    for tbl in pg.extract_tables():
                        for row in tbl:
                            celulas = [c.strip() if c else "" for c in row]
                            if any(celulas):
                                linhas.append(" | ".join(celulas))
                    t = pg.extract_text()
                    if t:
                        linhas.append(t)
                texto = "\n".join(linhas)
            else:
                # Demais documentos: texto corrido
                texto = "\n".join(t for pg in pdf.pages if (t := pg.extract_text()))
    except Exception as e:
        return False, f"=== {nome} — erro na extração: {e} ==="
    ok = len(texto.strip()) > 20
    return ok, f"=== {nome} ===\n" + texto


def _extrair_docx(nome: str, content: bytes) -> tuple[bool, str]:
    try:
        d = docx_lib.Document(io.BytesIO(content))
        linhas = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                celulas = [c.text.strip() for c in row.cells]
                if any(celulas):
                    linhas.append(" | ".join(celulas))
    except Exception as e:
        return False, f"=== {nome} — erro na extração: {e} ==="
    texto = "\n".join(linhas)
    return len(texto.strip()) > 20, f"=== {nome} ===\n" + texto


def _extrair_texto_puro(nome: str, content: bytes) -> tuple[bool, str]:
    for enc in ("utf-8", "latin-1"):
        try:
            texto = content.decode(enc)
            return len(texto.strip()) > 20, f"=== {nome} ===\n" + texto
        except UnicodeDecodeError:
            continue
    return False, f"=== {nome} — não foi possível decodificar como texto ==="


def _extrair_html(nome: str, content: bytes) -> tuple[bool, str]:
    ok, texto = _extrair_texto_puro(nome, content)
    return ok, re.sub(r"<[^>]+>", " ", texto)


def _extrair_bruto_best_effort(nome: str, content: bytes) -> tuple[bool, str]:
    """Fallback pra formato binário sem parser dedicado (.doc antigo, .rtf, etc.) —
    varre bytes decodificáveis e mantém só sequências imprimíveis de 4+ chars.
    SEMPRE conta como não-confiável (ok=False) — qualidade não garantida, nunca
    deve sozinho justificar seguir pra análise (ver ExtracaoInsuficiente)."""
    bruto = content.decode("latin-1", errors="ignore")
    pedacos = re.findall(r"[ -~À-ÿ]{4,}", bruto)
    texto = "\n".join(pedacos)
    return False, f"=== {nome} — extração best-effort (formato binário sem parser), conferir original ===\n" + texto


EXTRATORES = {
    ".pdf":  _extrair_pdf,
    ".docx": _extrair_docx,
    ".txt":  _extrair_texto_puro,
    ".csv":  _extrair_texto_puro,
    ".html": _extrair_html,
    ".htm":  _extrair_html,
}


def _extrair_arquivo(nome: str, content: bytes) -> tuple[bool, str]:
    ext = "." + nome.lower().rsplit(".", 1)[-1] if "." in nome else ""
    extrator = EXTRATORES.get(ext)
    if extrator:
        return extrator(nome, content)
    if b"%PDF" in content[:8]:
        return _extrair_pdf(nome, content)
    print(f"    ⚠ sem parser dedicado pra '{ext or '(sem extensão)'}' — usando extração best-effort", file=sys.stderr)
    return _extrair_bruto_best_effort(nome, content)


def baixar_documentos(cnpj: str, ano: int, seq: int) -> str:
    seq_s = f"{seq:06d}"
    try:
        arquivos = pncp_get(f"{cnpj}/compras/{ano}/{seq_s}/arquivos")
    except Exception as e:
        raise ExtracaoInsuficiente(f"Falha ao listar arquivos do PNCP: {e}")

    resultados = []  # lista de (ok: bool, texto: str) — todo arquivo, sem exceção
    for arq in arquivos:
        titulo = arq.get("titulo") or arq.get("nome") or "documento"
        url    = arq.get("url") or arq.get("uri") or ""
        if not url:
            continue
        print(f"  ↓ {titulo}", file=sys.stderr)
        try:
            r = requests.get(url, timeout=60)
        except Exception as e:
            print(f"    erro: {e}", file=sys.stderr)
            resultados.append((False, f"=== {titulo} — falha de download: {e} ==="))
            continue
        if r.status_code != 200:
            resultados.append((False, f"=== {titulo} — HTTP {r.status_code} ==="))
            continue
        if r.content[:4] == b"PK\x03\x04":
            # ZIP (ou .docx, que também é um zip por dentro) — abre e lê TODO
            # arquivo de dentro, independente do formato (PNCP mistura pdf/docx/etc).
            try:
                z = zipfile.ZipFile(io.BytesIO(r.content))
                nomes_zip = z.namelist()
            except zipfile.BadZipFile:
                print(f"    ZIP inválido: {titulo}", file=sys.stderr)
                resultados.append((False, f"=== {titulo} — ZIP inválido/corrompido ==="))
                continue
            # .docx é ele mesmo um zip (word/document.xml etc.) — se o titulo já
            # aponta pra um .docx, trata o download inteiro como esse arquivo,
            # não como container de outros documentos.
            if titulo.lower().endswith(".docx") and any(n.startswith("word/") for n in nomes_zip):
                resultados.append(_extrair_docx(titulo, r.content))
                continue
            for nome in nomes_zip:
                if nome.endswith("/"):
                    continue
                print(f"    └ {nome}", file=sys.stderr)
                resultados.append(_extrair_arquivo(nome, z.read(nome)))
        else:
            resultados.append(_extrair_arquivo(titulo, r.content))

    chars_confiaveis = sum(len(t) for ok, t in resultados if ok)
    n_ok             = sum(1 for ok, _ in resultados if ok)
    if n_ok == 0 or chars_confiaveis < LIMIAR_CHARS_CONFIAVEIS:
        raise ExtracaoInsuficiente(
            f"Extração insuficiente pra confiar: {n_ok} documento(s) lido(s) com sucesso, "
            f"{chars_confiaveis} chars confiáveis (mínimo {LIMIAR_CHARS_CONFIAVEIS}). "
            f"Analisar sem base documental real geraria alucinação — processo interrompido de propósito."
        )

    return "\n\n".join(t for _, t in resultados)

# ─── PROMPT ───────────────────────────────────────────────────────────────────

PROMPT_SISTEMA = """\
Você é especialista em licitações públicas brasileiras. Analise o edital e retorne SOMENTE JSON válido,\
 sem texto adicional, markdown ou blocos de código. Siga a estrutura exata abaixo.

Regras obrigatórias:
- Medidas de pneu SEMPRE no formato "NNN/NN RNN" (ex: "175/70 R14", "205/75 R16C"). Nunca use traço no lugar de barra.
- prazo_entrega e prazo_pagamento: copie o texto exato do documento (ex: "30 dias após recebimento do empenho"). \
Nunca retorne "conforme TR", "conforme edital" ou null — se não encontrar, procure mais.
- obs do item: use apenas para especificações relevantes (ex: "Nylon, 8 lonas", "AT 50%", "com câmara"). \
Deixe vazio se não houver informação adicional relevante.
- Câmara: "Com" se o item incluir câmara de ar, "Sem" caso contrário. Atenção a itens que listam câmara junto ao pneu.
- catmat: código numérico do CATMAT/SIASG se presente no documento. Se não encontrar, use "".
- habilitacao: máximo 80 chars por campo. Use abreviações: RFB+PGFN, FGTS, CNDT, Est.+Mun., doc. adm., CC/CCMEI. \
Se não houver exigência, escreva "Sem exigência". Não use frases completas.
"""

PROMPT_ESTRUTURA = """\
{
  "titulo": "ORGAO-UF | PE NNNNN/AAAA | X itens / Y UN | R$ VALOR | UASG XXXXXX",
  "data_sessao_iso": "AAAA-MM-DD",
  "cabecalho": {
    "objeto": "...",
    "modalidade": "Pregão Eletrônico",
    "lei": "14.133/2021",
    "valor_total": "R$ XX.XXX,XX",
    "data_sessao": "DD/MM/AAAA às HHh",
    "modo_disputa": "Aberto",
    "exclusivo_meepp": true,
    "srp": true,
    "prazo_entrega": "30 dias do empenho",
    "prazo_pagamento": "30 dias da NF"
  },
  "habilitacao": {
    "juridica": "Contrato Social/CCMEI · doc. de administradores",
    "fiscal": "CNPJ · RFB+PGFN · FGTS · CNDT · Fazenda Estadual",
    "economico_financeira": "Certidão negativa de falência — sem balanço",
    "tecnica": "Atestado de fornecimento (PJ pública ou privada)",
    "declaracoes": "Conjunta · ME-EPP · Concordância",
    "via_sicaf": true,
    "exige_balanco": false,
    "prazo_docs_externos": "2h"
  },
  "specs_comuns": "Novo · Radial · INMETRO · ...",
  "itens": [
    {
      "numero": 1,
      "catmat": "000000",
      "medida": "195/65 R15",
      "camara": "Sem",
      "qtde": 0,
      "valor_unit": "R$ 000,00",
      "valor_total": "R$ 00.000,00",
      "obs": ""
    }
  ],
  "entrega_local": "Endereço completo de entrega",
  "leilao": {
    "plataforma": "Comprasnet / PNCP",
    "modo": "Aberto",
    "criterio": "Menor preço por item",
    "prazo_recurso": "3 dias úteis",
    "prazo_arp": "5 dias da convocação"
  },
  "alertas": {
    "bloqueantes": [
      "Atestado Técnico — não possui · crítico para DD/MM",
      "INMETRO por item — solicitar laudo ao fornecedor"
    ],
    "atencao": [
      "Item X com câmara — confirmar INMETRO + Treadwear",
      "DOT ≤ 12 meses do empenho — risco estoque antigo"
    ],
    "vantagens": [
      "Exclusivo ME/EPP · sem balanço · atestado simples",
      "SICAF — docs externos: 2h para enviar na sessão"
    ]
  },
  "pontos_chave": {
    "favoraveis": "Exclusivo ME/EPP · sem balanço · SICAF · atestado simples",
    "atencao": "Item X câmara · DOT empenho · Frete ~Xkm",
    "competitivo": "3–6 licitantes · itens A+B+C = 75% do valor · Piso 50%"
  },
  "documentos_proposta": [
    {"documento": "Contrato Social / CCMEI", "status": "⏳ Aguarda JUCEC"},
    {"documento": "CNPJ (Cartão)", "status": "✅ Ativo"},
    {"documento": "CND Conjunta RFB/PGFN", "status": "✅ Válida até DD/MM/AAAA"},
    {"documento": "CRF — FGTS", "status": "⏳ Pendente"},
    {"documento": "CNDT — Certidão Trabalhista", "status": "✅ Válida até DD/MM/AAAA"},
    {"documento": "Regularidade Estadual (CE)", "status": "✅ Emitida"},
    {"documento": "Certidão Negativa de Falência", "status": "⏳ Pendente"},
    {"documento": "Atestado de Capacidade Técnica", "status": "❌ Não possui"},
    {"documento": "Laudo INMETRO por item", "status": "❌ Solicitar ao fornecedor"},
    {"documento": "Declarações (edital)", "status": "🔄 Gerar por pregão"}
  ]
}
"""

def analisar(texto: str, metadados: dict, itens_api: list) -> dict:
    cliente = anthropic.Anthropic()
    contexto = (
        f"METADADOS PNCP:\n{json.dumps(metadados, ensure_ascii=False)}\n\n"
        f"ITENS (API):\n{json.dumps(itens_api, ensure_ascii=False)}\n\n"
        f"TEXTO DOS DOCUMENTOS:\n{texto[:MAX_TEXT_CHARS]}"
    )
    resp = cliente.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=PROMPT_SISTEMA,
        messages=[{"role": "user", "content":
            f"Estrutura esperada:\n{PROMPT_ESTRUTURA}\n\nEDITAL:\n{contexto}"}]
    )
    raw = resp.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\n?", "", raw)
    raw = re.sub(r"\n?```$", "", raw)
    return json.loads(raw)

# ─── NOTION — HELPERS ─────────────────────────────────────────────────────────

def _notion_headers():
    token = os.environ.get("NOTION_TOKEN", "")
    if not token:
        raise RuntimeError("Variável NOTION_TOKEN não definida.")
    return {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Notion-Version": NOTION_VERSION,
    }

def notion_get_children(page_id: str):
    r = requests.get(f"{NOTION_BASE}/blocks/{page_id}/children?page_size=100",
                     headers=_notion_headers(), timeout=30)
    r.raise_for_status()
    return r.json().get("results", [])

def notion_delete(block_id: str):
    requests.delete(f"{NOTION_BASE}/blocks/{block_id}",
                    headers=_notion_headers(), timeout=30)

def notion_append(page_id: str, children: list):
    for i in range(0, len(children), 100):
        r = requests.patch(
            f"{NOTION_BASE}/blocks/{page_id}/children",
            headers=_notion_headers(),
            json={"children": children[i:i + 100]},
            timeout=30,
        )
        if not r.ok:
            print(f"Notion API error {r.status_code}: {r.text}", file=sys.stderr)
        r.raise_for_status()

def notion_update_props(page_id: str, titulo: str, data_iso: str):
    props = {
        "Nome do Processo": {"title": [{"text": {"content": titulo}}]}
    }
    if data_iso:
        props["Data da Sessão"] = {"date": {"start": data_iso}}
    r = requests.patch(f"{NOTION_BASE}/pages/{page_id}",
                       headers=_notion_headers(),
                       json={"properties": props}, timeout=30)
    r.raise_for_status()

# ─── NOTION — BUILDERS DE BLOCO ───────────────────────────────────────────────

def _rt(text: str, bold=False, italic=False):
    obj = {"type": "text", "text": {"content": str(text)}}
    if bold or italic:
        obj["annotations"] = {}
        if bold:
            obj["annotations"]["bold"] = True
        if italic:
            obj["annotations"]["italic"] = True
    return obj

def b_divider():
    return {"object": "block", "type": "divider", "divider": {}}

def b_h1(text: str):
    return {"object": "block", "type": "heading_1",
            "heading_1": {"rich_text": [_rt(text)]}}

def b_h2(text: str):
    return {"object": "block", "type": "heading_2",
            "heading_2": {"rich_text": [_rt(text)]}}

def b_para(text: str, bold=False, italic=False):
    return {"object": "block", "type": "paragraph",
            "paragraph": {"rich_text": [_rt(text, bold=bold, italic=italic)]}}

def b_callout(text: str, emoji="ℹ️"):
    return {"object": "block", "type": "callout",
            "callout": {"rich_text": [_rt(text)],
                        "icon": {"type": "emoji", "emoji": emoji}}}

def b_table(headers: list, rows: list):
    """Cria bloco table com header row."""
    def row(cells):
        return {"object": "block", "type": "table_row",
                "table_row": {"cells": [[_rt(str(c))] for c in cells]}}

    return {
        "object": "block",
        "type": "table",
        "table": {
            "table_width": len(headers),
            "has_column_header": True,
            "has_row_header": False,
            "children": [row(headers)] + [row(r) for r in rows],
        },
    }

# ─── BUILDER DO CARD ──────────────────────────────────────────────────────────

def build_blocks(a: dict) -> list:
    c    = a["cabecalho"]
    h    = a["habilitacao"]
    al   = a["alertas"]
    pk   = a["pontos_chave"]
    leil = a["leilao"]

    bl = []

    # ── ANÁLISE DE EDITAL ────────────────────────────────────────────────────
    bl.append(b_divider())
    bl.append(b_h1("📋 ANÁLISE DE EDITAL"))

    srp_s  = "SRP (1 ano, prorrogável)" if c.get("srp") else ""
    meepp  = "Exclusivo ME/EPP" if c.get("exclusivo_meepp") else ""
    linha3 = " · ".join(filter(None, [
        srp_s, meepp,
        f"Entrega: {c['prazo_entrega']}",
        f"Pagamento: {c['prazo_pagamento']}",
    ]))
    bl.append(b_para(f"{c['modalidade']} · {c['lei']} · {c['objeto'][:100]}"))
    bl.append(b_para(f"Sessão: {c['data_sessao']}  ·  Total: {c['valor_total']}  ·  {c['modo_disputa']} — Menor preço por item"))
    bl.append(b_para(linha3))

    # Tabela de alertas (sem cabeçalho visível)
    alert_rows = (
        [["🔴", x] for x in al["bloqueantes"]] +
        [["⚠️", x] for x in al["atencao"]] +
        [["✅", x] for x in al["vantagens"]]
    )
    bl.append(b_table(["", ""], alert_rows))

    # ── HABILITAÇÃO ──────────────────────────────────────────────────────────
    bl.append(b_divider())
    bl.append(b_h2("HABILITAÇÃO"))
    bl.append(b_table(
        ["Categoria", "Documentos"],
        [
            ["Jurídica",             h["juridica"]],
            ["Fiscal",               h["fiscal"]],
            ["Econômico-Financeira", h["economico_financeira"]],
            ["Técnica",              h["tecnica"]],
            ["Declarações",          h["declaracoes"]],
        ]
    ))
    sicaf = f"Via SICAF · docs externos: até {h.get('prazo_docs_externos','2h')}"
    if not h.get("exige_balanco"):
        sicaf += " · sem exigência de balanço"
    bl.append(b_callout(sicaf, "ℹ️"))

    # ── PRODUTOS ─────────────────────────────────────────────────────────────
    bl.append(b_divider())
    bl.append(b_h2("PRODUTOS"))
    if a.get("specs_comuns"):
        bl.append(b_para(f"Specs comuns: {a['specs_comuns']}"))

    itens = a["itens"]
    qtde_total = sum(i.get("qtde", 0) for i in itens)
    rows_itens = [
        [i["numero"], i.get("catmat", ""), i["medida"],
         i.get("camara", "—"), i["qtde"],
         i.get("valor_unit", ""), i.get("valor_total", ""), i.get("obs", "")]
        for i in itens
    ] + [["TOTAL", "", "", "", qtde_total, "", c["valor_total"], ""]]

    bl.append(b_table(
        ["Item", "CATMAT", "Medida", "Câmara", "Qtde", "Valor UN", "Valor Total", "Obs"],
        rows_itens
    ))
    if a.get("entrega_local"):
        bl.append(b_para(f"Entrega: {a['entrega_local']}"))

    # ── PONTOS-CHAVE ─────────────────────────────────────────────────────────
    bl.append(b_divider())
    bl.append(b_h2("PONTOS-CHAVE"))
    bl.append(b_table(
        ["Aspecto", "Detalhe"],
        [
            ["✅ Favoráveis",   pk["favoraveis"]],
            ["⚠️ Atenção",      pk["atencao"]],
            ["🏆 Competitivo",  pk["competitivo"]],
        ]
    ))

    # ── LEILÃO ───────────────────────────────────────────────────────────────
    bl.append(b_divider())
    bl.append(b_h2("LEILÃO"))
    bl.append(b_table(
        ["Campo", "Detalhe"],
        [
            ["Data/hora",            c["data_sessao"]],
            ["Plataforma",           leil["plataforma"]],
            ["Modo",                 leil["modo"]],
            ["Critério",             leil["criterio"]],
            ["Prazo recurso",        leil.get("prazo_recurso", "")],
            ["Prazo ARP assinatura", leil.get("prazo_arp", "")],
        ]
    ))

    # ── DOCUMENTOS DA PROPOSTA ───────────────────────────────────────────────
    bl.append(b_h2("DOCUMENTOS DA PROPOSTA"))
    docs = a.get("documentos_proposta", [])
    bl.append(b_table(
        ["Documento", "Status"],
        [[d["documento"], d["status"]] for d in docs]
    ))
    bl.append(b_para(
        "✅ = disponível · ⏳ = pendente · ❌ = não possui · 🔄 = gerar por pregão",
        italic=True
    ))

    return bl

# ─── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    dry_run   = "--dry-run" in sys.argv
    args      = [a for a in sys.argv[1:] if not a.startswith("--")]

    if len(args) < 1 or (not dry_run and len(args) < 2):
        print(__doc__, file=sys.stderr)
        sys.exit(1)

    pncp_arg  = args[0]
    notion_id = args[1].replace("-", "") if not dry_run else None

    # 1. PNCP
    print("📡 Conectando ao PNCP...", file=sys.stderr)
    cnpj, ano, seq = parse_pncp(pncp_arg)
    seq_s = f"{seq:06d}"
    print(f"   CNPJ {cnpj} · {ano} · {seq_s}", file=sys.stderr)

    metadados = pncp_get(f"{cnpj}/compras/{ano}/{seq_s}")
    itens_api = pncp_get(f"{cnpj}/compras/{ano}/{seq_s}/itens")

    # 2. Documentos
    print("📄 Baixando documentos...", file=sys.stderr)
    try:
        texto = baixar_documentos(cnpj, ano, seq)
    except ExtracaoInsuficiente as e:
        print(f"\n❌ ERRO — processo interrompido: {e}", file=sys.stderr)
        print("   Nenhuma chamada ao Claude foi feita. Nenhuma escrita no Notion foi feita.", file=sys.stderr)
        sys.exit(2)
    print(f"   {len(texto):,} chars extraídos de {pncp_arg}", file=sys.stderr)

    # 3. Claude
    print("🤖 Analisando com Claude...", file=sys.stderr)
    analise = analisar(texto, metadados, itens_api)
    print(f"   {len(analise['itens'])} itens · {analise['cabecalho']['valor_total']}", file=sys.stderr)

    # 4. Saída
    if dry_run:
        sys.stdout.buffer.write(json.dumps(analise, ensure_ascii=False, indent=2).encode("utf-8"))
        print("\n✅ Dry-run concluído — JSON impresso no stdout", file=sys.stderr)
        return

    print("✏️  Atualizando Notion...", file=sys.stderr)
    notion_update_props(notion_id, analise["titulo"], analise.get("data_sessao_iso", ""))

    children = notion_get_children(notion_id)
    for bloco in children:
        if bloco["type"] not in ("file", "pdf", "image"):
            notion_delete(bloco["id"])

    novos = build_blocks(analise)
    notion_append(notion_id, novos)
    print(f"   {len(novos)} blocos inseridos", file=sys.stderr)

    print(f"\n✅ https://app.notion.com/p/{notion_id}", file=sys.stderr)

    # Salva JSON para auditoria
    out = Path(f"analise_{cnpj}_{ano}_{seq_s}.json")
    out.write_text(json.dumps(analise, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"   JSON salvo em {out}", file=sys.stderr)


if __name__ == "__main__":
    from pathlib import Path
    main()
